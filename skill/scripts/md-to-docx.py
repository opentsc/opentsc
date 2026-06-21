#!/usr/bin/env python3
"""
Convert a markdown proposal to a formatted Word (.docx) document.

Usage:
    python md-to-docx.py input.md [output.docx]

Features:
- H1/H2/H3 headings with appropriate font sizes
- Bold inline markdown (**text**) parsed to Word bold runs
- Bullet lists with proper List Bullet style
- Tables (markdown pipe tables) rendered as Table Grid
- Centered title and metadata lines
- Chinese font support (微软雅黑)
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cn_font(run, size=None):
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if size:
        run.font.size = Pt(size)


def parse_inline(p, text):
    """Parse **bold** inline markdown into Word runs."""
    parts = re.split(r'\*\*(.*?)\*\*', text)
    for idx, part in enumerate(parts):
        run = p.add_run(part)
        set_cn_font(run)
        if idx % 2 == 1:
            run.bold = True


def convert(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    lines = md_text.split('\n')
    table_rows = []
    i = 0

    while i < len(lines):
        s = lines[i].strip()

        if not s:
            i += 1
            continue

        # Horizontal rule
        if s == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 50)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)
            i += 1
            continue

        # H1 (not H2+)
        if s.startswith('# ') and not s.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s[2:])
            run.bold = True
            set_cn_font(run, 18)
            i += 1
            continue

        # H3
        if s.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(s[4:])
            run.bold = True
            set_cn_font(run, 12)
            i += 1
            continue

        # H2
        if s.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(s[3:])
            run.bold = True
            set_cn_font(run, 14)
            i += 1
            continue

        # Table row
        if '|' in s and s.startswith('|'):
            cells = [c.strip() for c in s.split('|')[1:-1]]
            # Skip separator row (---, :---:)
            if all(set(c.strip()) <= set('- :') for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is also a table row
            if i + 1 < len(lines) and '|' in lines[i+1] and lines[i+1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table — create it
                if table_rows:
                    nc = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=nc)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ri, rd in enumerate(table_rows):
                        for ci, ct in enumerate(rd):
                            if ci < nc:
                                cell = table.cell(ri, ci)
                                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', ct)
                                cell.text = clean
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        set_cn_font(run, 10)
                                        if ri == 0:
                                            run.bold = True
                    doc.add_paragraph()
                table_rows = []
                i += 1
                continue

        # Bold metadata lines (centered)
        if s.startswith('**呈') or s.startswith('**提交') or s.startswith('**2026'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parse_inline(p, s)
            i += 1
            continue

        # Bullet
        if s.startswith('- '):
            text = s[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline(p, s)
        i += 1

    doc.save(docx_path)
    print(f'Saved: {docx_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python md-to-docx.py input.md [output.docx]')
        sys.exit(1)
    md_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else md_path.rsplit('.', 1)[0] + '.docx'
    convert(md_path, docx_path)
